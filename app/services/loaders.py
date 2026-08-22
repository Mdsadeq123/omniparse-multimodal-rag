from pathlib import Path

from langchain_core.documents import Document
from pypdf import PdfReader


def load_document(file_path: Path, doc_id: str, filename: str) -> list[Document]:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(file_path, doc_id, filename)
    if suffix == ".docx":
        return _load_docx(file_path, doc_id, filename)
    if suffix in {".txt", ".md", ".markdown"}:
        return _load_text(file_path, doc_id, filename)
    raise ValueError(f"Unsupported document type: {suffix}")


def _load_pdf(file_path: Path, doc_id: str, filename: str) -> list[Document]:
    reader = PdfReader(str(file_path))
    docs: list[Document] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "doc_id": doc_id,
                    "source": filename,
                    "type": "document",
                    "page": str(i + 1),
                },
            )
        )
    return docs


def _load_docx(file_path: Path, doc_id: str, filename: str) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        return []
    full_text = "\n\n".join(paragraphs)
    return [
        Document(
            page_content=full_text,
            metadata={
                "doc_id": doc_id,
                "source": filename,
                "type": "document",
                "page": "1",
            },
        )
    ]


def _load_text(file_path: Path, doc_id: str, filename: str) -> list[Document]:
    text = file_path.read_text(encoding="utf-8", errors="replace")
    if not text.strip():
        return []
    return [
        Document(
            page_content=text,
            metadata={
                "doc_id": doc_id,
                "source": filename,
                "type": "document",
                "page": "1",
            },
        )
    ]
